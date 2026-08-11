import os
import glob
import docx
from doc2docx import convert
from pypdf import PdfReader
from docxtpl import DocxTemplate
import streamlit as st
import PyPDF2
import pdfplumber
import pandas as pd
import json
import re
import time
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openai import OpenAI

# ==================== PAGE CONFIG ====================
st.set_page_config(
    page_title="Multi-Provider Talent CV & TOR Extractor",
    page_icon="📄",
    layout="wide"
)

HISTORY_FILE = "session_history.json"
KEYS_FILE = "saved_api_keys.json"
INPUT_DIR = "input_tor"
OUTPUT_DIR = "output_job_posts"
TEMPLATE_FILE = "template.docx"

os.makedirs(INPUT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ==================== PERSISTENT KEYS MANAGEMENT ====================
def load_saved_keys():
    if os.path.exists(KEYS_FILE):
        try:
            with open(KEYS_FILE, "r") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_keys_to_file(keys_dict):
    try:
        with open(KEYS_FILE, "w") as f:
            json.dump(keys_dict, f, indent=2)
    except Exception:
        pass

saved_file_keys = load_saved_keys()
api_providers = ["openrouter", "gemini", "groq", "mistral", "together", "cohere"]

for provider in api_providers:
    session_key = f"{provider}_key"
    param_val = st.query_params.get(session_key, "")
    file_val = saved_file_keys.get(session_key, "")
    if session_key not in st.session_state:
        st.session_state[session_key] = param_val or file_val or ""

# ==================== LOCAL HISTORY MANAGEMENT ====================
def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_history(new_results):
    history = load_history()
    updated = new_results + history
    updated = updated[:20]
    with open(HISTORY_FILE, "w") as f:
        json.dump(updated, f, indent=2)

# ==================== EXCEL FORMATTING FUNCTION ====================
def export_formatted_excel(dataframe, filename="Extracted_Candidates.xlsx"):
    dataframe.to_excel(filename, index=False, engine='openpyxl')
    
    wb = openpyxl.load_workbook(filename)
    ws = wb.active

    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    
    for col in ws.iter_cols(min_row=1, max_row=1):
        for cell in col:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

    col_widths = {
        'File Name': 20,
        'Name': 22,
        'Email': 25,
        'Phone': 20,
        'Designation': 25,
        'Current Organization': 25,
        'Previous Organizations': 35,
        'Experience Summary': 30,
        'Education': 35,
        'Source': 22
    }

    for col in ws.columns:
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        col_name = str(col[0].value)
        width = col_widths.get(col_name, 22)
        ws.column_dimensions[col_letter].width = width

        for cell in col[1:]:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(filename)
    return filename

# ==================== DETERMINISTIC (REGEX) & TOR PARSERS ====================
def extract_email(text):
    match = re.search(r'[a-zA-Z0-9%+\_.-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return match.group(0).strip() if match else "N/A"

def extract_phone(text):
    phones = re.findall(r'(?:\+?88)?01[3-9]\d{8}', text)
    if phones:
        return ", ".join(list(dict.fromkeys(phones)))
    return "N/A"

def extract_education_hardcode(text):
    """Parses education keywords directly from text/tables."""
    edu_matches = []
    
    keywords = ["JSC", "SSC", "HSC", "Dakhil", "Alim", "Fazil", "Kamil", "BSc", "BA", "BCom", "MBBS", "Diploma", "Masters"]
    
    lines = text.split('\n')
    for line in lines:
        for kw in keywords:
            if re.search(r'\b' + kw + r'\b', line, re.IGNORECASE):
                segment = line.strip()
                if len(segment) < 100 and segment not in edu_matches:
                    edu_matches.append(segment)
                    
    if edu_matches:
        return " | ".join(edu_matches[:3])
    return "Not Found"

def preprocess_files():
    """Automatically converts any legacy .doc files in the folder to .docx format on the fly."""
    doc_files = glob.glob(os.path.join(INPUT_DIR, "*.doc")) + glob.glob(os.path.join(INPUT_DIR, "*.DOC"))
    for file_path in doc_files:
        try:
            convert(file_path)
        except Exception:
            pass

def extract_text_from_file(file_path):
    ext = file_path.lower().split('.')[-1]
    full_text = ""
    
    if ext == "docx":
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            if p.text.strip():
                full_text += p.text.strip() + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                cell_texts = [cell.text.strip().replace('\n', ' ') for cell in row.cells if cell.text.strip()]
                if cell_texts:
                    full_text += " | ".join(cell_texts) + "\n"
                    
    elif ext == "pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text += text + "\n"
                
    return full_text

def parse_tor_dynamically(file_path):
    filename = os.path.basename(file_path)
    raw_text = extract_text_from_file(file_path)

    lines = [line.strip() for line in raw_text.split('\n') if line.strip()]
    
    clean_title = filename.rsplit(".", 1)[0].replace("_", " ")
    for line in lines:
        if "Position title" in line or "Position:" in line:
            parts = line.split("|")
            if len(parts) > 1:
                clean_title = parts[-1].strip()
                break

    responsibilities = []
    education = []
    experience = []
    languages = []
    
    current_section = None
    for line in lines:
        lower_line = line.lower()
        
        if "responsibilities and accountabilities" in lower_line or "iii. responsibilities" in lower_line or "job responsibility" in lower_line:
            current_section = "resp"
            continue
        elif "education and experience" in lower_line or "v. education and experience" in lower_line or "iv. required qualifications" in lower_line or "educational requirements" in lower_line:
            current_section = "edu_block"
            continue
        elif "experience requirements" in lower_line:
            current_section = "exp"
            continue
        elif "languages" in lower_line or "vi. languages" in lower_line or "v. languages" in lower_line:
            current_section = "lang"
            continue
        elif "vi. competencies" in lower_line or "notes" in lower_line or "iv. competencies" in lower_line or "apply procedure" in lower_line:
            current_section = None
            
        cleaned_line = line
        for tag in ["<td colspan=3/>", "<td colspan=2/>", "<td colspan=1/>", "<td/>"]:
            cleaned_line = cleaned_line.replace(tag, "")
        
        sub_parts = [p.strip(" |•-") for p in cleaned_line.split("|") if p.strip(" |•-")]
        
        for part in sub_parts:
            if not part or len(part) < 3:
                continue
            if part.lower() in ["education", "experience", "required", "advantageous", "either:", "or", "educational requirements", "experience requirements"]:
                continue

            if current_section == "resp":
                if part not in responsibilities:
                    responsibilities.append(part)
            elif current_section == "edu_block":
                if "year" in part.lower() or "diploma" in part.lower() or "degree" in part.lower() or "bachelor" in part.lower():
                    if part not in education:
                        education.append(part)
                else:
                    if part not in experience:
                        experience.append(part)
            elif current_section == "exp":
                if part not in experience:
                    experience.append(part)
            elif current_section == "lang":
                if part not in languages:
                    languages.append(part)

    responsibilities = list(dict.fromkeys(responsibilities))
    education = list(dict.fromkeys(education))
    experience = list(dict.fromkeys(experience))
    languages = list(dict.fromkeys(languages))

    if not responsibilities:
        responsibilities = ["Carry out regular field visits and monitor project activities."]
    if not education:
        education = ["Bachelor’s degree in relevant field or equivalent technical certification/diploma."]
    if not experience:
        experience = ["Relevant professional experience in project implementation and field operations."]
    if not languages:
        languages = ["Fluency in English and Bengali is required."]

    competencies = {
        "values_heading": "Values - all staff members must abide by and demonstrate these core values:",
        "values_list": [
            "Inclusion and respect for diversity: respects and promotes individual and cultural differences.",
            "Integrity and transparency: maintains high ethical standards and acts consistent with organizational principles.",
            "Professionalism: demonstrates ability to work in a composed, competent, and committed manner."
        ],
        "core_heading": "Core Competencies – behavioural indicators",
        "core_list": [
            "Teamwork: develops and promotes effective collaboration within and across units to achieve shared goals.",
            "Delivering results: produces and delivers quality results in a service-oriented and timely manner.",
            "Managing and sharing knowledge: continuously seeks to learn, share knowledge and innovate.",
            "Accountability: takes ownership for achieving priorities and assumes responsibility for own actions.",
            "Communication: encourages clear and open communication and explains complex matters clearly."
        ],
        "managerial_heading": "Managerial Competencies – behavioural indicators",
        "managerial_list": [
            "Leadership: provides a clear sense of direction and leads by example.",
            "Empowering others and building trust: creates an enabling environment where staff can contribute their best.",
            "Strategic thinking and vision: works strategically to realize organizational goals."
        ]
    }

    return {
        "job_title": clean_title,
        "responsibilities": responsibilities[:12],
        "education": education[:5],
        "experience": experience[:8],
        "competencies": competencies,
        "languages": languages[:4]
    }

# ==================== NAVIGATION TABS ====================
tab_main, tab_tor, tab_api, tab_history = st.tabs(["🚀 CV Extractor", "📋 TOR Job Post Generator", "⚙️ Persistent API Settings", "📜 Local Session History"])

# ==================== TAB 3: API SETTINGS ====================
with tab_api:
    st.header("🔑 Multi-Provider API Configuration")
    st.caption("Keys entered here are auto-saved locally so you DON'T have to re-enter them on refresh!")

    def update_key(provider_name):
        val = st.session_state[f"input_{provider_name}"]
        st.session_state[f"{provider_name}_key"] = val
        st.query_params[f"{provider_name}_key"] = val
        current_keys = load_saved_keys()
        current_keys[f"{provider_name}_key"] = val
        save_keys_to_file(current_keys)

    col_a, col_b = st.columns(2)
    
    with col_a:
        st.subheader("1. OpenRouter")
        st.markdown("[🔗 Get OpenRouter Key](https://openrouter.ai/keys)")
        st.text_input("OpenRouter Key", value=st.session_state["openrouter_key"], type="password", key="input_openrouter", on_change=update_key, args=("openrouter",))

        st.subheader("2. Google Gemini")
        st.markdown("[🔗 Get Gemini Key](https://aistudio.google.com/)")
        st.text_input("Gemini Key", value=st.session_state["gemini_key"], type="password", key="input_gemini", on_change=update_key, args=("gemini",))

        st.subheader("3. Groq AI")
        st.markdown("[🔗 Get Groq Key](https://console.groq.com/keys)")
        st.text_input("Groq Key", value=st.session_state["groq_key"], type="password", key="input_groq", on_change=update_key, args=("groq",))

    with col_b:
        st.subheader("4. Mistral AI")
        st.markdown("[🔗 Get Mistral Key](https://console.mistral.ai/)")
        st.text_input("Mistral Key", value=st.session_state["mistral_key"], type="password", key="input_mistral", on_change=update_key, args=("mistral",))

        st.subheader("5. Together AI")
        st.markdown("[🔗 Get Together AI Key](https://api.together.ai/)")
        st.text_input("Together AI Key", value=st.session_state["together_key"], type="password", key="input_together", on_change=update_key, args=("together",))

        st.subheader("6. Cohere AI")
        st.markdown("[🔗 Get Cohere Key](https://dashboard.cohere.com/api-keys)")
        st.text_input("Cohere Key", value=st.session_state["cohere_key"], type="password", key="input_cohere", on_change=update_key, args=("cohere",))

    st.markdown("---")
    st.subheader("🧪 Test API Connections")
    
    if st.button("Test Saved API Keys"):
        results_status = []
        if st.session_state["openrouter_key"]:
            try:
                c = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=st.session_state["openrouter_key"])
                c.chat.completions.create(model="openrouter/free", messages=[{"role":"user","content":"hi"}], max_tokens=5)
                results_status.append("✅ OpenRouter: Connected successfully!")
            except Exception as e:
                results_status.append(f"❌ OpenRouter Error: {str(e)[:60]}")

        if st.session_state["groq_key"]:
            try:
                c = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=st.session_state["groq_key"])
                c.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role":"user","content":"hi"}], max_tokens=5)
                results_status.append("✅ Groq: Connected successfully!")
            except Exception as e:
                results_status.append(f"❌ Groq Error: {str(e)[:60]}")

        if st.session_state["gemini_key"]:
            try:
                c = OpenAI(base_url="https://generativelanguage.googleapis.com/v1beta/openai/", api_key=st.session_state["gemini_key"])
                c.chat.completions.create(model="gemini-2.0-flash", messages=[{"role":"user","content":"hi"}], max_tokens=5)
                results_status.append("✅ Google Gemini: Connected successfully!")
            except Exception as e:
                results_status.append(f"❌ Gemini Error: {str(e)[:60]}")

        if not results_status:
            st.warning("Please enter at least one API key above.")
        else:
            for status in results_status:
                st.write(status)

# ==================== PDF EXTRACTION ENGINE ====================

def extract_text_from_pdf(file):
    text = ""
    file.seek(0)
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(layout=True) or page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip(), "pdfplumber"
    except Exception:
        pass
    
    try:
        file.seek(0)
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        if text.strip():
            return text.strip(), "PyPDF2"
    except Exception:
        pass
    
    return "", "none"

def clean_and_parse_json(text):
    text = text.strip()
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()
    try:
        return json.loads(text)
    except Exception:
        pass
    
    match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except Exception:
            pass

    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            pass
            
    return None

def run_ai_extraction(text):
    lines = [line for line in text.split('\n') if not line.strip().startswith('•')]
    compressed_text = "\n".join(lines)[:3500]

    prompt_system = """Extract candidate info from CV text into valid JSON.
JSON Keys Required:
- Name
- Designation
- Current Organization
- Previous Organizations (List all prior employers separated by '|'. Return "None Listed" if none exist)
- Experience Summary (Total years e.g., "25.7 yrs")
- Education (Degrees, Exams, and Institutes e.g. "SSC from Sarail Annada Govt. High School")

Return JSON ONLY."""

    user_prompt = f"CV TEXT:\n{compressed_text}"

    providers = []
    if st.session_state.get("openrouter_key"):
        providers.extend([
            ("OpenRouter", "https://openrouter.ai/api/v1", st.session_state["openrouter_key"], "openrouter/free"),
            ("OpenRouter", "https://openrouter.ai/api/v1", st.session_state["openrouter_key"], "meta-llama/llama-3.1-8b-instruct:free")
        ])
    if st.session_state.get("groq_key"):
        providers.append(("Groq", "https://api.groq.com/openai/v1", st.session_state["groq_key"], "llama-3.3-70b-versatile"))
    if st.session_state.get("gemini_key"):
        providers.append(("Gemini", "https://generativelanguage.googleapis.com/v1beta/openai/", st.session_state["gemini_key"], "gemini-2.0-flash"))

    if not providers:
        return None, "No API key configured."

    for provider_name, base_url, key, model in providers:
        try:
            client = OpenAI(base_url=base_url, api_key=key)
            res = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": prompt_system},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.0
            )
            raw = res.choices[0].message.content
            parsed = clean_and_parse_json(raw)
            if parsed:
                parsed["Source"] = f"{provider_name} ({model})"
                return parsed, None
        except Exception:
            time.sleep(1)
            continue

    return None, "All API providers timed out or hit rate limits."

# ==================== TAB 1: MAIN EXTRACTOR ====================
with tab_main:
    st.write("Upload PDF CVs below to extract candidate data into Excel.")
    uploaded_files = st.file_uploader("Drop CV PDFs here", type=["pdf"], accept_multiple_files=True)

    if st.button("🚀 Process CVs", type="primary") and uploaded_files:
        results = []
        progress = st.progress(0)
        
        for i, file in enumerate(uploaded_files):
            raw_text, _ = extract_text_from_pdf(file)
            
            if not raw_text.strip():
                st.error(f"Could not read {file.name}")
                continue
            
            email_val = extract_email(raw_text)
            phone_val = extract_phone(raw_text)
            fallback_edu = extract_education_hardcode(raw_text)

            data, err = run_ai_extraction(raw_text)
            
            if data:
                data["File Name"] = file.name
                data["Email"] = email_val
                data["Phone"] = phone_val if data.get("Phone") in ["N/A", "", None] else data.get("Phone")
                
                if not data.get("Education") or data.get("Education") in ["Not Found", "None Listed", ""]:
                    data["Education"] = fallback_edu

                results.append(data)
                st.success(f"Successfully processed {file.name}")
            else:
                st.warning(f"AI failed for {file.name}, using Regex Fallback.")
                results.append({
                    "File Name": file.name,
                    "Name": file.name.replace(".pdf", ""),
                    "Email": email_val,
                    "Phone": phone_val,
                    "Designation": "N/A",
                    "Current Organization": "N/A",
                    "Previous Organizations": "N/A",
                    "Experience Summary": "N/A",
                    "Education": fallback_edu,
                    "Source": "Python Regex Fallback"
                })
            
            progress.progress((i + 1) / len(uploaded_files))

        if results:
            st.balloons()
            save_history(results)
            
            df = pd.DataFrame(results)
            expected_cols = ['File Name', 'Name', 'Email', 'Phone', 'Designation', 'Current Organization', 'Previous Organizations', 'Experience Summary', 'Education', 'Source']
            df = df.reindex(columns=[c for c in expected_cols if c in df.columns])
            
            st.dataframe(df, use_container_width=True)
            
            tsv_data = df.to_csv(sep="\t", index=False)
            st.subheader("📋 Copy Results directly")
            st.code(tsv_data, language="text")

            excel_file = export_formatted_excel(df)
            with open(excel_file, "rb") as f:
                st.download_button(
                    label="📥 Download Structured Excel File",
                    data=f.read(),
                    file_name="Candidates_Summary.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

# ==================== TAB 2: TOR JOB POST GENERATOR ====================
with tab_tor:
    st.header("📋 TOR to Job Post Document Generator")
    st.markdown("Upload Terms of Reference (TOR) files in `.docx`, `.doc`, or `.pdf` format to automatically convert and structure them into polished job circulars using `template.docx`.")

    if not os.path.exists(TEMPLATE_FILE):
        st.error(f"⚠️ Missing required template file: `{TEMPLATE_FILE}`. Please place it in the application root directory.")
    
    uploaded_tors = st.file_uploader("Upload TOR Files", type=["docx", "doc", "pdf"], accept_multiple_files=True, key="tor_uploader")

    if st.button("⚙️ Generate Job Posts", type="primary") and uploaded_tors:
        # Save uploaded files into INPUT_DIR
        for uploaded_file in uploaded_tors:
            temp_path = os.path.join(INPUT_DIR, uploaded_file.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

        preprocess_files()

        tor_files = []
        for ext in ("*.docx", "*.DOCX", "*.pdf", "*.PDF"):
            tor_files.extend(glob.glob(os.path.join(INPUT_DIR, ext)))
        
        if not tor_files:
            st.warning(f"⚠️ No supported TOR files found in the '{INPUT_DIR}' folder!")
        else:
            progress_bar = st.progress(0)
            success_count = 0

            for idx, file_path in enumerate(tor_files):
                filename = os.path.basename(file_path)
                try:
                    template = DocxTemplate(TEMPLATE_FILE)
                    data = parse_tor_dynamically(file_path)

                    context = {
                        "job_title": data["job_title"],
                        "location": "Cox's Bazar",
                        "employment_status": "Contractual",
                        "deadline": "August 30, 2026",
                        "salary": "As per company policy",
                        
                        "responsibilities": data["responsibilities"],
                        "education": data["education"],
                        "experience": data["experience"],
                        "competencies": data["competencies"],
                        "languages": data["languages"]
                    }
                    
                    template.render(context)
                    
                    output_name = filename.rsplit(".", 1)[0] + ".docx"
                    output_path = os.path.join(OUTPUT_DIR, f"Generated_{output_name}")
                    template.save(output_path)
                    success_count += 1
                    st.success(f"✅ Generated: Generated_{output_name}")
                except Exception as e:
                    st.error(f"❌ Error processing {filename}: {e}")
                
                progress_bar.progress((idx + 1) / len(tor_files))

            if success_count > 0:
                st.balloons()
                st.info("All generated files are saved in the `output_job_posts` directory.")
                
                # Provide direct download buttons for generated files
                generated_results = glob.glob(os.path.join(OUTPUT_DIR, "*.docx"))
                for gen_file in generated_results:
                    with open(gen_file, "rb") as f:
                        st.download_button(
                            label=f"📥 Download {os.path.basename(gen_file)}",
                            data=f.read(),
                            file_name=os.path.basename(gen_file),
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=gen_file
                        )

# ==================== TAB 4: LOCAL HISTORY ====================
with tab_history:
    st.header("📜 Session History (Last 20 Extractions)")
    history_data = load_history()
    
    if history_data:
        df_hist = pd.DataFrame(history_data)
        st.dataframe(df_hist, use_container_width=True)
        
        tsv_hist = df_hist.to_csv(sep="\t", index=False)
        st.code(tsv_hist, language="text")
        
        if st.button("Clear Saved History"):
            if os.path.exists(HISTORY_FILE):
                os.remove(HISTORY_FILE)
                st.rerun()
    else:
        st.info("No saved local history found yet.")
